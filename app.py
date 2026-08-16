import json
import os
import shutil
from datetime import datetime, timezone, timedelta
import pandas as pd
import streamlit as st
import urllib.parse
from docx import Document
import re
import io
import streamlit.components.v1 as components

try:
    import cv2
    import numpy as np
    OPENCV_DISPONIVEL = True
except ImportError:
    OPENCV_DISPONIVEL = False

st.set_page_config(
    page_title="Premium Wines - Galpão",
    page_icon="imagem premium.jpeg",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
    <style>
    .stApp { background: linear-gradient(135deg, #F8F9FA 0%, #E9ECEF 100%); color: #1A1A1A; font-family: 'Poppins', sans-serif; overscroll-behavior-y: none; }
    [data-testid="stSidebar"] { display: none; }
    
    footer {visibility: hidden;}
    #MainMenu {visibility: hidden;}
    [data-testid="stStatusWidget"] {display: none;}
    
    label { color: #7A1C2E !important; font-weight: 700 !important; font-size: 0.95rem !important; }
    .wine-card { background-color: #FFFFFF; color: #1A1A1A; border-radius: 14px; padding: 16px; margin-bottom: 12px; border: 1px solid #E9ECEF; box-shadow: 0px 4px 12px rgba(0, 0, 0, 0.03); }
    .wine-title { color: #7A1C2E; font-size: 1.1rem; font-weight: 700; margin-bottom: 4px; }
    .badge-pallet-grande { background-color: #7A1C2E; color: #FFFFFF; padding: 6px 14px; border-radius: 8px; font-weight: 700; font-size: 1rem; display: inline-block; }
    .stButton button { background-color: #7A1C2E !important; color: #FFFFFF !important; border-radius: 12px !important; font-weight: 600 !important; border: none !important; padding: 10px 16px !important; width: 100%; white-space: pre-wrap; }
    </style>
""", unsafe_allow_html=True,
)

NOME_ARQUIVO = "estoque_galpao_pro.json"
ARQUIVO_USUARIOS = "usuarios_galpao.json"
ARQUIVO_LOGS = "logs_auditoria.json"
ARQUIVO_PEDIDOS = "pedidos_matriz.json"
PASTA_BACKUP = "backups_estoque"
PASTA_FOTOS = "fotos_vinhos"
SENHA_DEV = "1980"

if not os.path.exists(PASTA_BACKUP):
    os.makedirs(PASTA_BACKUP)
if not os.path.exists(PASTA_FOTOS):
    os.makedirs(PASTA_FOTOS)

LISTA_CORREDORES = [f"Corredor {i:02d}" for i in range(1, 26)]
LISTA_LOCAIS_TIPO = ["Pallet", "Prateleira"]
LISTA_NUMEROS_LOCAL = [f"Item {i:02d}" for i in range(1, 26)]
LISTA_LADOS = ["Direito", "Esquerdo", "Centro / Único"]
OPCOES_CAIXA = ["Caixa com 12 garrafas", "Caixa com 6 garrafas", "Caixa com 3 garrafas", "Caixa com 2 garrafas", "Garrafa Avulsa (1 un)", "Outra quantidade"]

def obter_horario_brasilia():
    fuso_brasilia = timezone(timedelta(hours=-3))
    return datetime.now(fuso_brasilia)

def obter_saudacao():
    hora = obter_horario_brasilia().hour
    if 0 <= hora < 12: return "Bom dia"
    elif 12 <= hora < 18: return "Boa tarde"
    else: return "Boa noite"

def realizar_backup(nome):
    if os.path.exists(nome):
        ts = obter_horario_brasilia().strftime("%Y%m%d_%H%M%S")
        shutil.copy(nome, os.path.join(PASTA_BACKUP, f"backup_{ts}_{nome}"))

def carregar_dados():
    estoque = []
    if os.path.exists(NOME_ARQUIVO):
        try:
            with open(NOME_ARQUIVO, "r", encoding="utf-8") as f: estoque = json.load(f)
        except: pass
    if not estoque:
        estoque = [{"nome": "Campana Merlot", "tipo": "Tinto", "safra": "2024", "localizacao": "Corredor 01 - Pallet Item 01", "lado": "Direito", "caixa": "Caixa com 12 garrafas", "codigo_barras": "7891008116632", "foto": ""}]
    return sorted(estoque, key=lambda x: x.get("nome", "").lower())

def salvar_dados(estoque):
    estoque_ordenado = sorted(estoque, key=lambda x: x.get("nome", "").lower())
    with open(NOME_ARQUIVO, "w", encoding="utf-8") as f: json.dump(estoque_ordenado, f, ensure_ascii=False, indent=4)
    realizar_backup(NOME_ARQUIVO)

def carregar_usuarios():
    if os.path.exists(ARQUIVO_USUARIOS):
        try:
            with open(ARQUIVO_USUARIOS, "r", encoding="utf-8") as f: return json.load(f)
        except: pass
    return [{"nome": "Vagner Souza", "cargo": "Administrador", "senha": "1980"}]

def salvar_usuarios(usuarios):
    with open(ARQUIVO_USUARIOS, "w", encoding="utf-8") as f: json.dump(usuarios, f, ensure_ascii=False, indent=4)

def carregar_logs():
    if os.path.exists(ARQUIVO_LOGS):
        try:
            with open(ARQUIVO_LOGS, "r", encoding="utf-8") as f: return json.load(f)
        except: pass
    return []

def registrar_log(usuario, acao, detalhes):
    logs = carregar_logs()
    logs.insert(0, {"data_hora": obter_horario_brasilia().strftime("%d/%m/%Y %H:%M:%S"), "usuario": usuario, "acao": acao, "detalhes": detalhes})
    with open(ARQUIVO_LOGS, "w", encoding="utf-8") as f: json.dump(logs, f, ensure_ascii=False, indent=4)

def carregar_pedidos():
    if os.path.exists(ARQUIVO_PEDIDOS):
        try:
            with open(ARQUIVO_PEDIDOS, "r", encoding="utf-8") as f: return json.load(f)
        except: pass
    return []

def salvar_pedidos(pedidos):
    with open(ARQUIVO_PEDIDOS, "w", encoding="utf-8") as f: json.dump(pedidos, f, ensure_ascii=False, indent=4)

def gerar_qr_code_api(texto):
    return f"https://api.qrserver.com/v1/create-qr-code/?size=250x250&data={urllib.parse.quote(texto)}"

def interpretar_linha_pedido(texto_linha):
    texto = texto_linha.strip()
    safra = ""
    quantidade = 1
    
    anos = re.findall(r'\b(20\d{2})\b', texto)
    if anos:
        safra = anos[0]
        texto_limpo = texto.replace(safra, "")
    else:
        texto_limpo = texto

    match_qtd = re.search(r'(?:/|\bcaixas?|\bqt[d]?\.?)\s*(\d+)', texto_limpo, re.IGNORECASE)
    if match_qtd:
        quantidade = int(match_qtd.group(1))
        texto_limpo = texto_limpo.replace(match_qtd.group(0), "")
    else:
        numeros_soltos = re.findall(r'\b(\d+)\b', texto_limpo)
        if numeros_soltos and numeros_soltos[-1] != safra:
            quantidade = int(numeros_soltos[-1])
            texto_limpo = texto_limpo.replace(numeros_soltos[-1], "")

    texto_limpo = re.sub(r'\bcaixas?\b', '', texto_limpo, flags=re.IGNORECASE)
    nome = re.sub(r'[/\|\-\–]+', '', texto_limpo).strip().title()
    return {"nome": nome, "safra": safra, "quantidade": quantidade, "separado": False, "qtd_separada": 0}

def extrair_pedidos_de_arquivo(arq):
    itens = []
    ext = arq.name.split('.')[-1].lower()
    try:
        if ext in ['xlsx', 'xls']:
            df = pd.read_excel(arq)
            for _, row in df.iterrows():
                nome_bruto = str(row.get('Nome', row.iloc[0] if len(row) > 0 else '')).strip()
                if nome_bruto and nome_bruto != 'Nan':
                    safra_col = str(row.get('Safra', row.iloc[1] if len(row) > 1 else '')).strip()
                    qtd_col = row.get('Quantidade', row.iloc[2] if len(row) > 2 else 1)
                    try: qtd = int(qtd_col) if pd.notnull(qtd_col) else 1
                    except: qtd = 1
                    
                    itens.append({
                        "nome": nome_bruto.title(), 
                        "safra": safra_col if safra_col != 'Nan' else '', 
                        "quantidade": qtd, 
                        "separado": False, 
                        "qtd_separada": 0
                    })
        elif ext == 'txt':
            linhas = [l.strip() for l in arq.getvalue().decode("utf-8").split("\n") if l.strip()]
            for l in linhas:
                itens.append(interpretar_linha_pedido(l))
    except: pass
    return itens

def extrair_linhas_de_arquivo(arq):
    linhas = []
    ext = arq.name.split('.')[-1].lower()
    try:
        if ext in ['xlsx', 'xls']:
            df = pd.read_excel(arq)
            for c in df.columns:
                for v in df[c].dropna():
                    if str(v).strip(): linhas.append(str(v).strip().title())
        elif ext == 'docx':
            doc = Document(arq)
            for p in doc.paragraphs:
                if p.text.strip(): linhas.append(p.text.strip().title())
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.strip(): linhas.append(cell.text.strip().title())
        elif ext == 'txt':
            linhas = [l.strip().title() for l in arq.getvalue().decode("utf-8").split("\n") if l.strip()]
    except: pass
    return linhas

def extrair_numero_corredor(localizacao):
    nums = re.findall(r'\d+', localizacao)
    if nums:
        return int(nums[0])
    return 999

def criar_documento_word_pedido(pedido):
    doc = Document()
    doc.add_heading(f"Relatório de Conferência de Pedido", level=1)
    doc.add_paragraph(f"Identificação: {pedido['id']}")
    doc.add_paragraph(f"Data de Criação: {pedido['data']}")
    doc.add_paragraph(f"Status Atual: {pedido.get('status', 'Pendente')}")
    doc.add_heading("Itens do Pedido:", level=2)
    
    for i, item in enumerate(pedido['itens'], 1):
        status_txt = "SEPARADO" if item.get('separado') else "PENDENTE"
        p_text = f"{i}. {item['nome']} (Safra: {item.get('safra', 'N/A')}) - Qtd Pedida: {item['quantidade']} | Qtd Separada: {item.get('qtd_separada', 0)} [{status_txt}]"
        doc.add_paragraph(p_text)
        
    f_stream = io.BytesIO()
    doc.save(f_stream)
    f_stream.seek(0)
    return f_stream

def componente_leitor_barcode(chave_sessao):
    html_code = f"""
    <div style="text-align: center;">
        <div id="reader_{chave_sessao}" style="width: 100%; max-width: 400px; margin: auto; border-radius: 12px; overflow: hidden;"></div>
        <p id="resultado_{chave_sessao}" style="font-weight: bold; color: #7A1C2E; margin-top: 8px;"></p>
    </div>
    <script src="https://unpkg.com/html5-qrcode"></script>
    <script>
      function onScanSuccess(decodedText, decodedResult) {{
        document.getElementById("resultado_{chave_sessao}").innerText = "Lido com sucesso: " + decodedText;
        const url = new URL(window.parent.location.href);
        url.searchParams.set('scanned_{chave_sessao}', decodedText);
        window.parent.history.replaceState({{}}, '', url);
        if (window.html5QrCode_{chave_sessao}) {{
            window.html5QrCode_{chave_sessao}.stop().catch(err => {{}});
        }}
      }}
      try {{
          const html5QrCode = new Html5Qrcode("reader_{chave_sessao}");
          window.html5QrCode_{chave_sessao} = html5QrCode;
          html5QrCode.start({{ facingMode: "environment" }}, {{ fps: 10, qrbox: {{ width: 250, height: 150 }} }}, onScanSuccess).catch(err => {{}});
      }} catch (e) {{}}
    </script>
    """
    components.html(html_code, height=320)

if "usuarios" not in st.session_state:
    st.session_state.usuarios = carregar_usuarios()

st.session_state.estoque = carregar_dados()
st.session_state.pedidos = carregar_pedidos()

if "menu_atual" not in st.session_state: st.session_state.menu_atual = "🏠 Home"
if "termo_busca" not in st.session_state: st.session_state.termo_busca = ""
if "codigo_capturado_cadastro" not in st.session_state: st.session_state.codigo_capturado_cadastro = ""
if "codigos_bipados_conferencia" not in st.session_state: st.session_state.codigos_bipados_conferencia = {}

qp = st.query_params

for key, val in list(qp.items()):
    if key.startswith("scanned_"):
        sess_key = key.replace("scanned_", "")
        valor_limpo = str(val).strip()
        
        if sess_key == "cad_barcode":
            st.session_state.codigo_capturado_cadastro = valor_limpo
        elif sess_key == "scanner_geral":
            st.session_state.termo_busca = valor_limpo
            st.session_state.menu_atual = "Filtros"
        else:
            st.session_state.codigos_bipados_conferencia[sess_key] = valor_limpo
            
        del st.query_params[key]
        st.rerun()

user_url = qp.get("user", None)
cargo_url = qp.get("cargo", "Operador")

if "usuario_logado" not in st.session_state or st.session_state.usuario_logado is None:
    if user_url:
        st.session_state.usuario_logado = {"nome": user_url, "cargo": cargo_url}
    else:
        st.session_state.usuario_logado = None

if st.session_state.usuario_logado is None:
    st.write("")
    _, cc, _ = st.columns([1, 1.3, 1])
    with cc:
        if os.path.exists("imagem premium.jpeg"):
            _, ci, _ = st.columns([1, 1.8, 1])
            with ci: st.image("imagem premium.jpeg", width=190)
        st.markdown("<h1 style='text-align: center; color: #7A1C2E; font-size: 1.6rem;'>PREMIUM WINES GALPÃO</h1>", unsafe_allow_html=True)
        
        tab1, tab2, tab3 = st.tabs(["🔑 Entrar", "👤 Criar Conta", "⚙️ Dev"])
        with tab1:
            with st.form("l_form"):
                u = st.text_input("Usuário").strip().title()
                p = st.text_input("Senha", type="password").strip()
                if st.form_submit_button("ENTRAR", use_container_width=True):
                    user = next((x for x in st.session_state.usuarios if x['nome'].lower() == u.lower() and x['senha'] == p), None)
                    if user:
                        st.session_state.usuario_logado = user
                        st.query_params["user"] = user['nome']
                        st.query_params["cargo"] = user.get('cargo', 'Operador')
                        st.rerun()
                    else: st.error("Dados incorretos.")
        with tab2:
            with st.form("c_form"):
                n = st.text_input("Nome").strip().title()
                s = st.text_input("Senha", type="password").strip()
                if st.form_submit_button("CADASTRAR", use_container_width=True):
                    if n and s:
                        novo = {"nome": n, "cargo": "Operador", "senha": s}
                        st.session_state.usuarios.append(novo)
                        salvar_usuarios(st.session_state.usuarios)
                        st.session_state.usuario_logado = novo
                        st.query_params["user"] = novo['nome']
                        st.query_params["cargo"] = novo['cargo']
                        st.rerun()
                    else: st.error("Preencha tudo.")
        with tab3:
            with st.form("d_form"):
                sp = st.text_input("Senha Mestra", type="password")
                if st.form_submit_button("DEV", use_container_width=True):
                    if sp == SENHA_DEV:
                        dev_user = {"nome": "Dev", "cargo": "Desenvolvedor"}
                        st.session_state.usuario_logado = dev_user
                        st.query_params["user"] = "Dev"
                        st.query_params["cargo"] = "Desenvolvedor"
                        st.rerun()
                    else: st.error("Senha incorreta.")
    st.stop()

ct1, ct2, ct3 = st.columns([3, 2, 1])
with ct1: st.markdown(f"🍷 <b>PREMIUM WINES</b> | Usuário: {st.session_state.usuario_logado['nome']} ({st.session_state.usuario_logado.get('cargo', 'Operador')})", unsafe_allow_html=True)
with ct2:
    if st.session_state.menu_atual != "🏠 Home":
        if st.button("⬅️ Voltar ao Menu", use_container_width=True): st.session_state.menu_atual = "🏠 Home"; st.rerun()
with ct3:
    if st.button("🚪 Sair", use_container_width=True):
        st.session_state.usuario_logado = None
        st.query_params.clear()
        st.session_state.menu_atual = "🏠 Home"
        st.rerun()

st.markdown("---")

if st.session_state.menu_atual == "🏠 Home":
    st.markdown(f"<p style='text-align: center; color: #666; margin-bottom: 0px;'>{obter_saudacao()},</p>", unsafe_allow_html=True)
    st.markdown(f"<h1 style='text-align: center; color: #7A1C2E; margin-top: 0px;'>{st.session_state.usuario_logado['nome']}! 👋</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #444; font-size: 0.95rem; margin-bottom: 25px;'>Escolha abaixo a opção desejada para gerenciar o galpão:</p>", unsafe_allow_html=True)
    
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("📦 Lista de Pedidos (Matriz)", use_container_width=True): st.session_state.menu_atual = "PedidosMatriz"; st.rerun()
    with c2:
        if st.button("🔍 Buscar / Filtros", use_container_width=True): st.session_state.menu_atual = "Filtros"; st.rerun()
    with c3:
        if st.button("🗺️ Mapa de Separação", use_container_width=True): st.session_state.menu_atual = "MapaSeparacao"; st.rerun()
    
    st.write("")
    c4, c5, c6 = st.columns(3)
    with c4:
        if st.button("🍷 Estoque Completo", use_container_width=True): st.session_state.menu_atual = "Estoque"; st.rerun()
    with c5:
        if st.button("➕ Cadastrar Vinho", use_container_width=True): st.session_state.menu_atual = "Cadastrar"; st.rerun()
    with c6:
        if st.button("📱 Gerar QR Code", use_container_width=True): st.session_state.menu_atual = "GerarQR"; st.rerun()
        
    st.write("")
    c7, c8, c9 = st.columns(3)
    with c7:
        if st.button("📷 Escanear Local", use_container_width=True): st.session_state.menu_atual = "Scanner"; st.rerun()
    with c8:
        if st.button("✏️ Editar Vinho", use_container_width=True): st.session_state.menu_atual = "Editar"; st.rerun()
    with c9:
        if st.button("📋 Histórico", use_container_width=True): st.session_state.menu_atual = "Historico"; st.rerun()
        
    if st.session_state.usuario_logado.get('cargo') == "Desenvolvedor":
        st.write("")
        if st.button("⚙️ Gerenciar Contas", use_container_width=True):
            st.session_state.menu_atual = "GerenciarUsuarios"
            st.rerun()

elif st.session_state.menu_atual == "PedidosMatriz":
    st.subheader("📦 Gerenciamento de Pedidos e Conferência Antierros")
    
    aba_ped1, aba_ped2 = st.tabs(["📋 Enviar/Novo Pedido", "🔍 Roteiro e Leitor de Caixa (Antierros)"])
    
    with aba_ped1:
        st.markdown("Envie a lista enviada pela matriz (Excel ou TXT) ou digite livremente.")
        st.info("ℹ️ Os pedidos salvos ficam guardados automaticamente no arquivo **pedidos_matriz.json** na pasta do sistema no computador.")
        
        with st.form("form_novo_pedido"):
            id_pedido = st.text_input("Identificação do Pedido / Loja", value=f"Pedido #{obter_horario_brasilia().strftime('%d/%m %H:%M')}")
            arq_pedido = st.file_uploader("Arquivo de Pedido (Excel ou TXT)", type=["xlsx", "xls", "txt"])
            texto_manual_pedido = st.text_area("Ex: Campana Merlot 2024 /05 Caixas", placeholder="Campana Merlot 2024 / 05 Caixas")
            
            cadastrar_clicado = st.form_submit_button("💾 Salvar Pedido no Sistema")
            
            if cadastrar_clicado:
                itens_novos = []
                if arq_pedido is not None:
                    itens_novos = extrair_pedidos_de_arquivo(arq_pedido)
                if texto_manual_pedido.strip():
                    for linha in texto_manual_pedido.split("\n"):
                        if linha.strip():
                            itens_novos.append(interpretar_linha_pedido(linha))
                
                if itens_novos:
                    novo_registro_pedido = {
                        "id": id_pedido,
                        "data": obter_horario_brasilia().strftime("%d/%m/%Y %H:%M"),
                        "itens": itens_novos,
                        "status": "Pendente"
                    }
                    st.session_state.pedidos.append(novo_registro_pedido)
                    salvar_pedidos(st.session_state.pedidos)
                    registrar_log(st.session_state.usuario_logado['nome'], "Novo Pedido Matriz", id_pedido)
                    st.success("Pedido salvo com sucesso no computador!")
                    st.rerun()
                else:
                    st.error("Adicione itens por arquivo ou texto.")
        
        if st.session_state.pedidos:
            st.markdown("---")
            st.markdown("### 📥 Baixar Pedido no Aparelho (Celular ou Computador)")
            st.markdown("Se você estiver abrindo pelo **celular** ou quiser uma cópia direta no seu dispositivo, baixe o arquivo JSON abaixo:")
            
            dados_json_str = json.dumps(st.session_state.pedidos, ensure_ascii=False, indent=4)
            st.download_button(
                label="📥 Baixar arquivo de pedidos (.json)",
                data=dados_json_str,
                file_name=f"pedidos_matriz_{obter_horario_brasilia().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json",
                use_container_width=True
            )
                    
    with aba_ped2:
        if not st.session_state.pedidos:
            st.info("Nenhum pedido cadastrado.")
        else:
            pedidos_nomes = [f"{p['id']} ({p['data']}) - Status: {p.get('status', 'Pendente')}" for p in st.session_state.pedidos]
            escolha_ped_str = st.selectbox("Selecione o Pedido:", pedidos_nomes)
            idx_ped = pedidos_nomes.index(escolha_ped_str)
            pedido_atual = st.session_state.pedidos[idx_ped]
            
            col_info_ped, col_del_ped = st.columns([3, 1])
            with col_info_ped:
                st.markdown(f"### Pedido: {pedido_atual['id']}")
            with col_del_ped:
                if st.button("🗑️ Excluir Pedido", use_container_width=True):
                    id_removido = pedido_atual['id']
                    st.session_state.pedidos.pop(idx_ped)
                    salvar_pedidos(st.session_state.pedidos)
                    registrar_log(st.session_state.usuario_logado['nome'], "Excluir Pedido", id_removido)
                    st.success("Pedido excluído com sucesso!")
                    st.rerun()
            
            st.markdown("---")
            docx_stream = criar_documento_word_pedido(pedido_atual)
            st.download_button(
                label="🖨️ Baixar Pedido em Word (.docx) para Imprimir ou Enviar",
                data=docx_stream,
                file_name=f"pedido_{re.sub(r'[^a-zA-Z0-9]', '_', pedido_atual['id'])}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
            ordem_separacao = st.radio("Direção da Rota pelos Corredores:", ["Crescente (Corredor 01 ao 25)", "Decrescente (Corredor 25 ao 01)"], horizontal=True)
            reversa = True if "Decrescente" in ordem_separacao else False
            
            itens_com_local = []
            for item in pedido_atual['itens']:
                nome_pedido_limpo = item['nome'].lower()
                safra_pedido = str(item.get('safra', '')).strip()
                
                vinho_encontrado = None
                for v in st.session_state.estoque:
                    v_nome = v.get('nome', '').lower()
                    v_safra = str(v.get('safra', '')).strip()
                    if v_nome in nome_pedido_limpo or nome_pedido_limpo in v_nome:
                        if safra_pedido and v_safra:
                            if safra_pedido == v_safra:
                                vinho_encontrado = v
                                break
                        else:
                            vinho_encontrado = v
                            break
                if not vinho_encontrado:
                    vinho_encontrado = next((v for v in st.session_state.estoque if v.get('nome', '').lower() in nome_pedido_limpo), None)
                
                loc = vinho_encontrado.get('localizacao', 'Corredor 01 - Pallet Item 01') if vinho_encontrado else 'Não cadastrado'
                corr_num = extrair_numero_corredor(loc)
                
                itens_com_local.append({
                    "item_original": item,
                    "vinho_estoque": vinho_encontrado,
                    "localizacao": loc,
                    "corredor_num": corr_num
                })
                
            itens_ordenados = sorted(itens_com_local, key=lambda x: x['corredor_num'], reverse=reversa)
            
            st.markdown("---")
            st.markdown("#### 🎯 Roteiro e Bipe de Caixas (Antierros)")
            
            todos_separados = True
            for i, obj in enumerate(itens_ordenados):
                item = obj['item_original']
                vinho_est = obj['vinho_estoque']
                loc = obj['localizacao']
                
                status_cor = "🟢" if item.get('separado') else "⏳"
                if not item.get('separado'): todos_separados = False
                
                key_bip_state = f"bip_val_{idx_ped}_{i}"
                if key_bip_state not in st.session_state.codigos_bipados_conferencia:
                    st.session_state.codigos_bipados_conferencia[key_bip_state] = ""

                with st.expander(f"{status_cor} {item['nome']} {item.get('safra', '')} | Qtd: {item['quantidade']} | 📍 {loc}"):
                    if vinho_est:
                        codigo_cadastrado_sistema = str(vinho_est.get('codigo_barras', '')).strip()
                        
                        st.markdown(f"**Galpão:** Safra: **{vinho_est.get('safra')}** | Local: **{vinho_est.get('localizacao')}**")
                        st.markdown(f"C. Barras Oficial Cadastrado: `{codigo_cadastrado_sistema if codigo_cadastrado_sistema else 'NÃO CADASTRADO'}`")
                        
                        st.markdown("---")
                        st.markdown("📷 **1º Passo: Bipe o código de barras da caixa**")
                        bip_caixa = st.text_input("Código de barras da caixa:", value=st.session_state.codigos_bipados_conferencia[key_bip_state], key=f"bip_txt_{idx_ped}_{i}")
                        if bip_caixa != st.session_state.codigos_bipados_conferencia[key_bip_state]:
                            st.session_state.codigos_bipados_conferencia[key_bip_state] = bip_caixa

                        componente_leitor_barcode(key_bip_state)

                        st.markdown("---")
                        st.markdown("📦 **2º Passo: Confirme a quantidade e finalize o item**")
                        qtd_informada = st.number_input("Quantidade que está levando:", min_value=1, value=item.get('quantidade', 1), key=f"qtd_inf_{idx_ped}_{i}")

                        pode_avancar = True
                        if not codigo_cadastrado_sistema:
                            st.warning("⚠️ Este vinho não possui código de barras cadastrado no sistema. A conferência será feita apenas por validação visual/manual.")
                        elif bip_caixa.strip() != codigo_cadastrado_sistema:
                            pode_avancar = False
                            if bip_caixa.strip():
                                st.error(f"❌ **ERRO DE CONFERÊNCIA!** O código bipado (`{bip_caixa}`) NÃO corresponde ao cadastro deste vinho (`{codigo_cadastrado_sistema}`).")

                        if st.button(f"✅ Confirmar Item #{i+1}", key=f"btn_sep_{idx_ped}_{i}"):
                            if not pode_avancar:
                                st.error("❌ Impossível confirmar: O código de barras bipado está incorreto!")
                            else:
                                item['separado'] = True
                                item['qtd_separada'] = qtd_informada
                                salvar_pedidos(st.session_state.pedidos)
                                registrar_log(st.session_state.usuario_logado['nome'], "Separar Item Pedido", f"{item['nome']}")
                                st.success("Item confirmado com sucesso!")
                                st.rerun()
                    else:
                        st.error("❌ Vinho não encontrado no estoque.")

            st.markdown("---")
            if todos_separados:
                st.success("🎉 Pedido 100% separado e conferido!")
                
                if st.button("💾 Finalizar e Arquivar Pedido", use_container_width=True, key="btn_finalizar_pedido_ok"):
                    pedido_atual['status'] = "Concluído"
                    salvar_pedidos(st.session_state.pedidos)
                    registrar_log(st.session_state.usuario_logado['nome'], "Concluir Pedido", pedido_atual['id'])
                    st.success("Pedido finalizado e salvo com sucesso!")
                    st.rerun()

elif st.session_state.menu_atual == "Filtros":
    st.subheader("🔍 Busca por Local ou Nome")
    termo_digitado = st.text_input("Filtrar por nome ou corredor/pallet:", value=st.session_state.termo_busca)
    tp = termo_digitado.strip().lower()
    if tp:
        res = [v for v in st.session_state.estoque if tp in v.get("nome", "").lower() or tp in v.get("localizacao", "").lower() or tp in str(v.get("codigo_barras", "")).lower()]
        if res:
            for v in sorted(res, key=lambda x: x.get("nome", "").lower()):
                st.markdown(f"<div class='wine-card'><div class='wine-title'>🍷 {v.get('nome')} ({v.get('safra')})</div><p>Tipo: <b>{v.get('tipo', 'N/A')}</b><br>C. Barras: <code>{v.get('codigo_barras', 'S/N')}</code><br><span class='badge-pallet-grande'>📍 {v.get('localizacao')}</span></p></div>", unsafe_allow_html=True)
        else:
            st.warning("Nenhum vinho encontrado.")

elif st.session_state.menu_atual == "MapaSeparacao":
    st.subheader("🗺️ Mapa de Separação")
    arq = st.file_uploader("Envie arquivo", type=["xlsx", "xls", "txt"])
    txt_man_input = st.text_area("Ou cole a lista manualmente:")
    if st.button("Gerar Rota"):
        linhas = extrair_linhas_de_arquivo(arq) if arq else []
        if txt_man_input.strip():
            linhas.extend([l.strip().title() for l in txt_man_input.split("\n") if l.strip()])
        if linhas:
            encontrados = [v for v in st.session_state.estoque if any(l.lower() in v.get("nome", "").lower() for l in linhas)]
            for v in encontrados:
                st.markdown(f"<div class='wine-card'><div class='wine-title'>🍷 {v.get('nome')} ({v.get('safra')})</div><span class='badge-pallet-grande'>📍 {v.get('localizacao')}</span></div>", unsafe_allow_html=True)

elif st.session_state.menu_atual == "Scanner":
    st.subheader("📷 Escanear Código de Barras ou QR Code")
    
    aba_scan1, aba_scan2 = st.tabs(["📱 Leitor com Câmera", "🖼️ Enviar Foto com QR Code"])
    
    with aba_scan1:
        st.markdown("Aponte a câmera para o código de barras da caixa ou QR Code do local:")
        componente_leitor_barcode("scanner_geral")
        
        if st.session_state.get("termo_busca"):
            termo = st.session_state.termo_busca
            st.info(f"🔍 Resultado da leitura: **{termo}**")
            
            resultados = [
                v for v in st.session_state.estoque 
                if termo.lower() in str(v.get("codigo_barras", "")).lower() 
                or termo.lower() in str(v.get("localizacao", "")).lower()
                or termo.lower() in str(v.get("nome", "")).lower()
            ]
            
            if resultados:
                st.success(f"Encontrado(s) {len(resultados)} registro(s)!")
                for v in resultados:
                    st.markdown(
                        f"<div class='wine-card'>"
                        f"<div class='wine-title'>🍷 {v.get('nome')} ({v.get('safra')})</div>"
                        f"<p>Tipo: <b>{v.get('tipo', 'N/A')}</b><br>"
                        f"C. Barras: <code>{v.get('codigo_barras', 'Não cadastrado')}</code><br>"
                        f"<span class='badge-pallet-grande'>📍 {v.get('localizacao')}</span></p>"
                        f"</div>", 
                        unsafe_allow_html=True
                    )
            else:
                st.warning("⚠️ Nenhum vinho ou local encontrado com este código exato.")
                
            if st.button("Limpar Busca"):
                st.session_state.termo_busca = ""
                st.rerun()

    with aba_scan2:
        foto = st.camera_input("Tirar foto de um QR Code impresso")
        if foto and OPENCV_DISPONIVEL:
            img = cv2.imdecode(np.frombuffer(foto.getvalue(), np.uint8), cv2.IMREAD_COLOR)
            detector = cv2.QRCodeDetector()
            val, _, _ = detector.detectAndDecode(img)
            if val:
                st.success(f"QR Code Lido: **{val}**")
                st.session_state.termo_busca = val
                st.session_state.menu_atual = "Filtros"
                st.rerun()
            else:
                st.warning("Nenhum QR Code detectado na imagem. Tente novamente.")
        elif foto and not OPENCV_DISPONIVEL:
            st.error("A biblioteca OpenCV não está disponível neste ambiente para decodificar a foto.")

elif st.session_state.menu_atual == "Estoque":
    st.subheader("🍷 Estoque Completo do Galpão")
    if st.session_state.estoque:
        for v in st.session_state.estoque:
            st.markdown(
                f"<div class='wine-card'>"
                f"<div class='wine-title'>🍷 {v.get('nome')} ({v.get('safra')})</div>"
                f"<p>Tipo: <b>{v.get('tipo', 'N/A')}</b><br>"
                f"C. Barras: <code>{v.get('codigo_barras', 'S/N')}</code><br>"
                f"<span class='badge-pallet-grande'>📍 {v.get('localizacao')}</span></p>"
                f"</div>",
                unsafe_allow_html=True
            )
    else:
        st.info("O estoque está vazio.")

elif st.session_state.menu_atual == "Cadastrar":
    st.subheader("➕ Cadastrar Novo Vinho")
    with st.form("form_cadastrar_vinho"):
        nome_cad = st.text_input("Nome do Vinho").strip().title()
        tipo_cad = st.selectbox("Tipo", ["Tinto", "Branco", "Rosé", "Espumante", "Fortificado"])
        safra_cad = st.text_input("Safra").strip()
        
        col_loc1, col_loc2, col_loc3, col_loc4 = st.columns(4)
        with col_loc1: cor = st.selectbox("Corredor", LISTA_CORREDORES)
        with col_loc2: tipo_l = st.selectbox("Tipo Local", LISTA_LOCAIS_TIPO)
        with col_loc3: num_l = st.selectbox("Número", LISTA_NUMEROS_LOCAL)
        with col_loc4: lado = st.selectbox("Lado", LISTA_LADOS)
        
        localizacao_completa = f"{cor} - {tipo_l} {num_l} ({lado})"
        
        caixa_cad = st.selectbox("Tipo de Caixa", OPCOES_CAIXA)
        codigo_cad = st.text_input("Código de Barras", value=st.session_state.codigo_capturado_cadastro).strip()
        
        if st.form_submit_button("Salvar Vinho"):
            if nome_cad:
                novo_vinho = {
                    "nome": nome_cad,
                    "tipo": tipo_cad,
                    "safra": safra_cad,
                    "localizacao": localizacao_completa,
                    "lado": lado,
                    "caixa": caixa_cad,
                    "codigo_barras": codigo_cad,
                    "foto": ""
                }
                st.session_state.estoque.append(novo_vinho)
                salvar_dados(st.session_state.estoque)
                registrar_log(st.session_state.usuario_logado['nome'], "Cadastrar Vinho", nome_cad)
                st.session_state.codigo_capturado_cadastro = ""
                st.success("Vinho cadastrado com sucesso!")
                st.rerun()
            else:
                st.error("O nome do vinho é obrigatório.")

elif st.session_state.menu_atual == "GerarQR":
    st.subheader("📱 Gerar QR Code de Localização")
    loc_qr = st.selectbox("Selecione o Corredor/Local:", [f"{c} - {t} {n}" for c in LISTA_CORREDORES for t in LISTA_LOCAIS_TIPO for n in LISTA_NUMEROS_LOCAL])
    url_qr = gerar_qr_code_api(loc_qr)
    st.image(url_qr, width=250, caption=f"QR Code: {loc_qr}")
    st.markdown(f"Link direto para impressão: [Abrir QR Code]({url_qr})", unsafe_allow_html=True)

elif st.session_state.menu_atual == "Editar":
    st.subheader("✏️ Editar ou Excluir Vinho")
    if not st.session_state.estoque:
        st.info("Nenhum vinho cadastrado para editar.")
    else:
        nomes_vinhos = [f"{v['nome']} ({v.get('safra', 'S/Safra')}) - {v.get('localizacao', '')}" for v in st.session_state.estoque]
        escolha_ed = st.selectbox("Selecione o Vinho:", nomes_vinhos)
        idx_vinho = nomes_vinhos.index(escolha_ed)
        v_atual = st.session_state.estoque[idx_vinho]
        
        with st.form("form_editar_vinho"):
            n_nome = st.text_input("Nome", value=v_atual.get('nome', '')).strip().title()
            n_safra = st.text_input("Safra", value=v_atual.get('safra', '')).strip()
            n_loc = st.text_input("Localização", value=v_atual.get('localizacao', '')).strip()
            n_barras = st.text_input("Código de Barras", value=v_atual.get('codigo_barras', '')).strip()
            
            col_b1, col_b2 = st.columns(2)
            with col_b1:
                btn_salvar = st.form_submit_button("💾 Salvar Alterações")
            with col_b2:
                btn_excluir = st.form_submit_button("🗑️ Excluir Vinho")
                
            if btn_salvar:
                v_atual['nome'] = n_nome
                v_atual['safra'] = n_safra
                v_atual['localizacao'] = n_loc
                v_atual['codigo_barras'] = n_barras
                salvar_dados(st.session_state.estoque)
                registrar_log(st.session_state.usuario_logado['nome'], "Editar Vinho", n_nome)
                st.success("Alterações salvas com sucesso!")
                st.rerun()
            elif btn_excluir:
                removido = st.session_state.estoque.pop(idx_vinho)
                salvar_dados(st.session_state.estoque)
                registrar_log(st.session_state.usuario_logado['nome'], "Excluir Vinho", removido.get('nome'))
                st.success("Vinho excluído com sucesso!")
                st.rerun()

elif st.session_state.menu_atual == "Historico":
    st.subheader("📋 Histórico de Auditoria do Galpão")
    logs = carregar_logs()
    if logs:
        for l in logs:
            st.markdown(f"**[{l['data_hora']}] {l['usuario']}** - *{l['acao']}*: {l['detalhes']}")
    else:
        st.info("Nenhum registro de log encontrado.")

elif st.session_state.menu_atual == "GerenciarUsuarios":
    st.subheader("⚙️ Gerenciamento de Contas e Usuários")
    if st.session_state.usuario_logado.get('cargo') != "Desenvolvedor":
        st.error("Acesso negado. Apenas desenvolvedores.")
    else:
        usuarios = st.session_state.usuarios
        for i, usr in enumerate(usuarios):
            st.markdown(f"**Usuário:** {usr['nome']} | **Cargo:** {usr.get('cargo', 'Operador')}")
